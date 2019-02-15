from docx import Document
import json
import os
import sys
from collections import OrderedDict

from libs.constants import Constants

class LoadData:
    '''
    class to load all data for program
    '''
    def __init__(self):
        ## Get directory for file path
        self.file_dir = os.path.dirname(__file__)
        self.head, self.tail = os.path.split(self.file_dir)
        if(self.tail != 'Gen-Desc.docx'): 
          self.file_dir = self.head
        pass

    def getFile(self, file_path):
        return os.path.join(self.file_dir, file_path)

    def getGenDescFile(self):
        return self.getFile("assests/Gen-Desc.docx")

    def getOutputFile(self, name = ""):
        return self.getFile(self.getOutPutFile(name))
    
    def getTemplateDocumentValues(self):
        return json.load(open(self.getFile("assests/document_value_template.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getTemplateTableValues(self):
        return json.load(open(self.getFile("assests/document_table_template.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)
    
    def getInputValues(self):
        self.path = self.getInputFilePath()
        return json.load(open(self.getFile(self.path), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getDocxDocument(self, design_code):
        if design_code == Constants.ASCE_710_LRFD:
            return Document(self.getFile("assests/Gen-Desc.docx"))
        elif design_code == Constants.ASCE_710_ASD:
            return Document(self.getFile("assests/Gen-Desc_710_asd.docx"))

    def getTable27_3_1(self):
        return json.load(open(self.getFile("assests/asce/table_27_3_1.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getWindDesignDefaults(self):
        return json.load(open(self.getFile("assests/wind_design_defaults.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getWindCoeffiecients(self):
        return json.load(open(self.getFile("assests/wind_coeffiecients.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getWindMapDefaults(self):
        return json.load(open(self.getFile("assests/windmap_defaults.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getInputFilePath(self):
        # if called with no arguments, call app data pick file from there
        path = None
        if (len(sys.argv) > 1):
            path = self.getFile(sys.argv[1])
            
        else:
            path = self.getFile(os.getenv('LOCALAPPDATA') + "\\Trevexs SSC\\data\\sample1.trsc")
        
        return path

    def getOutPutFile(self, name):
        try: 
            head = os.path.split(self.path)[0]
            file_name = "Gen-Desc-%s.DOC" %name
            return os.path.join(head, file_name)

        except AttributeError:
            # Messages.showError("There is no data to use to generate output file")
            print("An error occured in getOutPutFile method")

    def getWindMapImagePath(self):
        try: 
            head = os.path.split(self.path)[0]
            file_name = "WINDMAP.JPG"
            return os.path.join(head, file_name)

        except AttributeError:
            # Messages.showError("There is no data to use to generate output file")
            print("An error occured in getOutPutFile method")