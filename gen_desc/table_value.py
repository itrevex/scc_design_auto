from docx.shared import Pt
from docx.enum.text import WD_BREAK

from libs.constants import Constants
from libs.general_methods import GeneralMethods
from .common import Common
from .run_properties import RunProperties
from gen_desc.wind_design.wind_design import WindDesign
from gen_desc.wind_design.constants import WindDesignConsts

class TableValue:
    def __init__(self, app_data, document, new_input_values):
        self.app_data = app_data
        self.document = document
        self.new_input_values = new_input_values
        self.template_table_values = app_data.getTemplateTableValues()
        self.calculateParameters()
        self.parapet_load = new_input_values[WindDesignConsts.PARAPET_LOAD]
        self.roof_angle = new_input_values[WindDesignConsts.ROOF_ANGLE]
        self.wind_design = WindDesign(app_data, self.getWindDesignProps())
        
        pass

    def getWindDesignProps(self):
        wind_design_props = {}
        wind_design_props["parapet_load"] = self.parapet_load
        wind_design_props["roof_angle"] = self.roof_angle
        wind_design_props["wind_load"] = self.wind_unit_load_kn

        return wind_design_props

    def getIdenfierParagraph(self, identifier_text):
        '''
        String identifier_text
        Find the paragraph that contains the idenfier text
        and return it
        '''
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if identifier_text in paragraph.text:
                            return paragraph
        return None

    def updateTableValues(self):
        for key, value in self.template_table_values.items():
            #key descriptive key of what to change
            identifier_text = value[Constants.TEMPLATE_IDENTIFIER]
            paragraph = self.getIdenfierParagraph(identifier_text)

            if (paragraph != None):
                paragraph.text = ""
                value_parts = value[Constants.VALUE_PARTS]
                self.updateValueParts(paragraph, value_parts)

                Common.giveFeedBack(key)
    
    def updateValueParts(self, paragraph, parts):
        for part_name, part in parts.items():
            text = self.getInputText(part, part_name, paragraph)
            run_props = RunProperties(text, part)
            TableValue.add_paragraph_run(paragraph, run_props)

    def getInputText(self, part, value, paragraph):

        input_from = TableValue.getParameter(part, Constants.INPUT_FROM)

        if (input_from == Constants.INPUT_FROM_INPUT):
            return self.new_input_values[value]
        elif(input_from == Constants.INPUT_FROM_CALCULATION):
            return self.calculatedValue(value, paragraph)

        return value

    def calculatedValue(self, value, paragraph):
        if(value == Constants.WIND_SPEED_MS):
            return self.speed_text

        if (value == Constants.KZ_VALUE):
            return self.kz_text

        if (value == Constants.WIND_UNIT_LOAD):
            return self.wind_unit_load_text

        if (value == Constants.WIND_UNIT_LOAD_KN):
            return self.wind_unit_load_kn_text

        if (value == Constants.WIND_DESIGN_X):
            runs = self.wind_design.wind_x.runs
            for run in runs:
                TableValue.add_paragraph_run(paragraph, run)
            pass

        if (value == Constants.WIND_DESIGN_Y):
            runs = self.wind_design.wind_y.runs
            for run in runs:
                TableValue.add_paragraph_run(paragraph, run)
            pass

        if (value == Constants.WIND_DESIGN_CALC_X):
            runs = self.wind_design.wind_calc_x.runs
            for run in runs:
                TableValue.add_paragraph_run(paragraph, run)
            pass

        if (value == Constants.WIND_DESIGN_CALC_Y):
            runs = self.wind_design.wind_calc_y.runs
            for run in runs:
                TableValue.add_paragraph_run(paragraph, run)
            pass

        return ""

    def calcKzValue(self):
        height_above_ground = self.new_input_values[Constants.HEIGHT_ABOVE_GROUND]
        self.kz = GeneralMethods(self.app_data).kzValue(float(height_above_ground))
        self.kz_text = GeneralMethods.kzValueText(self.kz)

    def calcSpeedMs(self):
        #calcualted wind speed in m/s and return value
        #first get wind_speed value from input
        wind_speed = self.new_input_values[Constants.WIND_SPEED]
        self.speed_ms, self.speed_text = GeneralMethods.windSpeedMPerSecond(wind_speed)

    def calcWindUnitLoad(self):
        self.wind_unit_load = GeneralMethods.calcWindUnitLoad(self.speed_ms, self.kz)
        self.wind_unit_load_text = GeneralMethods.windUnitLoadText(self.wind_unit_load)

    def calcWindUnitLoadKn(self):
        self.wind_unit_load_kn =  self.wind_unit_load/1000.
        self.wind_unit_load_kn_text = GeneralMethods.textFourPlaces(self.wind_unit_load_kn)

    def calculateParameters(self):
        self.calcSpeedMs()
        self.calcKzValue()
        self.calcWindUnitLoad()
        self.calcWindUnitLoadKn()

    @staticmethod
    def getParameter(part, key):
        try:
            return part[key]
        except KeyError:
            return ""
            
    @staticmethod    
    def add_paragraph_run(paragraph, run_props):
        run = paragraph.add_run(run_props.text)
        run = run_props.applyRunProps(run)

    def trials(self):
        self.wind_design.trials()