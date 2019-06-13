'''
    class is called by gen_desc class
    the class is entered through updateTableValues() method 
    after the constructor call
'''

from docx.shared import Pt
from docx.enum.text import WD_BREAK

from libs.constants import Constants
from libs.general_methods import GeneralMethods
from .common import Common
from .run_properties import RunProperties
from wind_design.wind_loads import WindLoads

class TableValue(WindLoads):
    WIND_LOAD = "wind_load"
    PARAPET_LOAD = "parapet_load"
    ROOF_ANGLE = "roof_angle"
    LENGTH_X = "length_x"
    LENGTH_Y = "length_y"
    HEIGHT = "height"

    def __init__(self, app_data, document, wind_design):
        
        super().__init__(app_data, wind_design)
        self.app_data = app_data
        self.document = document
        self.new_input_values = wind_design.props
        self.template_table_values = app_data.getTemplateTableValues()
        self.calculateParameters()
        self.wind_design = wind_design #WindDesign(app_data, new_input_values, self.wind_unit_load_kn)
        
        pass

    def getWindDesignProps(self):
        wind_design_props = {}
        return wind_design_props

    def getIdenfierParagraph(self, identifier_text):
        '''
        String identifier_text
        Find the paragraph that contains the idenfier text
        and return it
        '''
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if identifier_text in paragraph.text:
                            return paragraph
        return None

    def updateTableValues(self):
        '''
        iterate through all items in document_table_template
        for each item find the key

            "table_dead_load":{ //key
                //value
                "identifier": "Roof Dead Loads = 0.10 kN/m2", //Constants.TEMPLATE_IDENTIFIER
                "value_parts": { //Constants.VALUE_PARTS
                    //key is the text in document to change, 
                    //the value give the properties that should apply to that text
                    for example underline, bold, text size
                    "Roof Dead Loads = ":{}, 
                    "roof_dead_load":{
                        "input_from": "input"
                    },
                    " kN/m":{},
                    "2":{
                        "superscript": "true"
                    }
                }
            }
        '''
        for key, value in self.template_table_values.items():
            #key descriptive key of what to change
            identifier_text = value[Constants.TEMPLATE_IDENTIFIER]
            paragraph = self.getIdenfierParagraph(identifier_text)

            if (paragraph != None):
                paragraph.text = ""
                value_parts = value[Constants.VALUE_PARTS]
                self.updateValueParts(paragraph, value_parts)

                Common.giveFeedBack(key)
    
    def updateValueParts(self, paragraph, parts):
        for part_name, part in parts.items():
            text = self.getInputText(part, part_name, paragraph)
            run_props = RunProperties(text, part)
            TableValue.add_paragraph_run(paragraph, run_props)

    def getInputText(self, part, value, paragraph):

        input_from = TableValue.getParameter(part, Constants.INPUT_FROM)

        if (input_from == Constants.INPUT_FROM_INPUT):
            return self.new_input_values[value]
        elif(input_from == Constants.INPUT_FROM_CALCULATION):
            return self.calculatedValue(value, paragraph, part)

        return value

    def calculatedValue(self, value, paragraph, part={}):
        if(value == Constants.WIND_SPEED_MS):
            return self.speed_text

        if (value == Constants.KZ_VALUE):
            return self.kz_text

        if (value == Constants.WIND_UNIT_LOAD):
            return self.wind_unit_load_text

        if (value == Constants.WIND_UNIT_LOAD_KN):
            return self.wind_unit_load_kn_text

        if (value == Constants.WIND_DESIGN_X):
            runs = self.wind_design.wind_x.runs
            for run in runs:
                TableValue.add_paragraph_run(paragraph, run)
            pass

        if (value == Constants.WIND_DESIGN_Y):
            runs = self.wind_design.wind_y.runs
            for run in runs:
                TableValue.add_paragraph_run(paragraph, run)
            pass

        if (value == Constants.WIND_DESIGN_CALC_X):
            runs = self.wind_design.wind_calc_x.runs
            for run in runs:
                TableValue.add_paragraph_run(paragraph, run)
            pass

        if (value == Constants.WIND_DESIGN_CALC_Y):
            runs = self.wind_design.wind_calc_y.runs
            for run in runs:
                TableValue.add_paragraph_run(paragraph, run)
            pass

        if (value == Constants.ROOF_ENCLOSURE):
            enclosure_type = self.new_input_values[Constants.ROOF_ENCLOSURE]
            if (enclosure_type == Constants.OPEN_ROOF):
                return "Enclosure Specification is Open"
            else:
                return "Enclosure Specification is Enclosed"

        if (value == Constants.INTERNAL_PRESSURE_VALUE):
            enclosure_type = self.new_input_values[Constants.ROOF_ENCLOSURE]
            if (enclosure_type == Constants.OPEN_ROOF):
                return '0.00'
            else:
                return '0.18'
            pass  
        
        if (value == Constants.SEISMIC):
            prop = part[Constants.SEISMIC_TABLE_PROPERTY]
            if prop == Constants.SEISMIC_SS:
                return "{:.3f}".format(float(self.new_input_values[value][prop])/100)
            return self.new_input_values[value][prop]

        # todo implement sms values if (value == ) 
        if (value == Constants.SEISMIC_SS):
            prop = Constants.SEISMIC
            self.ss = float(self.new_input_values[prop][value])/100
            return "{:.3f}".format(self.ss) 

        if (value == Constants.SEISMIC_FA):
            prop = Constants.SEISMIC
            self.fa = float(self.new_input_values[prop][value])
            return "{:.2f}".format(self.fa)
        
        if (value == Constants.SEISMIC_SMS_VALUE):
            self.sms = self.fa * self.ss
            return "{:.4f}".format(self.sms)

        if (value == Constants.SEISMIC_SDS_VALUE):
            self.sds = (2/3) * self.sms
            return "{:.4f}".format(self.sds)

        if (value == Constants.SEISMIC_IMPORTANCE_FACTOR):
            prop = Constants.SEISMIC
            self.seismic_importance = float(self.new_input_values[prop][value])
            return "{:.2f}".format(self.seismic_importance)
            
        if (value == Constants.SEISMIC_CS_VALUE):
            self.cs = self.sds / (6/self.seismic_importance)
            return "{:.4f}".format(self.cs)

        if (value == Constants.SEISMIC_CS_MIN_VALUE):
            self.cs_min = 0.044 * self.sds * self.seismic_importance
            return "{:.4f}".format(self.cs_min)

        return ""

    def calcKzValue(self):
        height_above_ground = self.new_input_values[Constants.HEIGHT_ABOVE_GROUND]
        self.kz = GeneralMethods(self.app_data).kzValue(float(height_above_ground))
        self.kz_text = GeneralMethods.kzValueText(self.kz)

    def calcSpeedMs(self):
        #calcualted wind speed in m/s and return value
        #first get wind_speed value from input
        wind_speed = self.new_input_values[Constants.WIND_SPEED]
        self.speed_ms, self.speed_text = GeneralMethods.windSpeedMPerSecond(wind_speed)

    def calcWindUnitLoad(self):
        self.wind_unit_load = GeneralMethods.calcWindUnitLoad(self.speed_ms, self.kz)
        self.wind_unit_load_text = GeneralMethods.windUnitLoadText(self.wind_unit_load)

    def calcWindUnitLoadKn(self):
        self.wind_unit_load_kn =  self.wind_unit_load/1000.
        self.wind_unit_load_kn_text = GeneralMethods.textFourPlaces(self.wind_unit_load_kn)

    def calculateParameters(self):
        self.calcSpeedMs()
        self.calcKzValue()
        self.calcWindUnitLoad()
        self.calcWindUnitLoadKn()


    @staticmethod
    def getParameter(part, key):
        try:
            return part[key]
        except KeyError:
            return ""
            
    @staticmethod    
    def add_paragraph_run(paragraph, run_props):
        run = paragraph.add_run(run_props.text)
        run_props.applyRunProps(run)

    def trials(self):
        self.wind_design.trials()