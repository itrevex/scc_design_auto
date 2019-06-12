from docx.shared import Pt

from libs.constants import Constants
from .common import Common

class DocValue:
    def __init__(self, app_data, document, new_input_values):
        self.document = document
        self.new_input_values = new_input_values
        self.template_doc_values = app_data.getTemplateDocumentValues()
        pass


    def getIdenfierParagraph(self, identifier_text):
        '''
        String identifier_text
        Find the paragraph that contains the idenfier text
        and return it
        '''
        for paragraph in self.document.paragraphs:
            if identifier_text in paragraph.text:
                return paragraph
        return None

    def getDocValueReplacementText(self, key):
        '''
        Get replacement value for key provided within the 
        input values provided by the user
        '''
        return self.new_input_values[key]

    def getNewParagraphText(self, paragraph, original_text, new_text):
        '''
        Replace original text with new text with the paragraph
        Change the paragh text the the changed value
        '''
        return paragraph.text.replace(original_text, new_text)

    def updateDocumentValues(self):
        for key, value in self.template_doc_values.items():
            if key == Constants.SEISMIC:
                try:
                    self.updateSeismicData(key, value)
                except KeyError:
                    print()
                    print ("Warning: No seismic data provided . . .")
                    print()
            else:
                new_text = self.new_input_values[key]
                self.updateDocValue(key, value, new_text)
            
    def updateDocValue(self, key, value, new_text):
        identifier_text = value[Constants.TEMPLATE_IDENTIFIER]
        
        #loop through paragraph to get identifier to save memory
        for paragraph in self.document.paragraphs:
            if identifier_text in paragraph.text:
                # paragraph = self.getIdenfierParagraph(identifier_text)
            
                if (self.replaceWholeParagraph(value) == Constants.TRUE):
                    paragraph_text  = new_text
                
                #add else if to allow change of paragraph in parts with
                #Runproperties too, don't break the old code
                else:
                    replace_text = self.getReplacementText(value)
                    paragraph_text  = paragraph.text.replace(replace_text, new_text)

                paragraph.text = ""
                paragraph_text_run = paragraph.add_run(paragraph_text)
                paragraph_text_run.font.size = Pt(12)
                Common.giveFeedBack(key)
        # self.giveFeedBack(key)
        
    def updateSeismicData(self, key, value):
        for seismic_key, seismic_value in value.items():
            new_text = self.new_input_values[key][seismic_key]
            self.updateDocValue(key, seismic_value, new_text)
            

    def replaceWholeParagraph(self, value):
        try:
            return value[Constants.REPLACE_ENTIRE_PARAGRAPH]
        except KeyError:
            return ""

    def getReplacementText(self, value):
        try:
            return value[Constants.SEISMIC_REPLACE_VALUE]
        except KeyError:
            return value[Constants.TEMPLATE_IDENTIFIER]