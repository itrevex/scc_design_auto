from docx import Document
import json, os, sys
from collections import OrderedDict

from .constants import Constants

class LoadData:
    '''
    class to load all data for program
    '''
    def __init__(self):
        ## Get directory for file path
        self.file_dir = os.path.dirname(__file__)
        self.head, self.tail = os.path.split(self.file_dir)
        if(self.tail != 'Gen-Desc.docx'): 
          self.file_dir = self.head
        self.path = self.getInputFilePath()

    def getFile(self, file_path):
        return os.path.join(self.file_dir, file_path)

    def getGenDescFile(self):
        return self.getFile("assests/Gen-Desc.docx")

    def getOutputFile(self, name = ""):
        return self.getFile(self.getOutPutFilePath(name))
    
    def getTemplateDocumentValues(self):
        return json.load(open(self.getFile("assests/templates/program/document_value_template.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getTemplateTableValues(self):
        return json.load(open(self.getFile("assests/templates/program/document_table_template.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)
    
    def getInputValues(self):
        return json.load(open(self.getFile(self.path), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getDocxDocument(self, design_code, enclosure='open', seismic=False):
        if seismic == True:
            return self.loadSeismicFiles(design_code, enclosure)
        else:
            return self.loadNonSeismicFiles(design_code, enclosure)

        
    def getFormGrsFile(self, design_code, enclosure='open'):
        if design_code == Constants.ASCE_710_LRFD:
            if enclosure == Constants.ENCLOSED_ROOF:
                return self.getFile("assests/templates/design/FORM_GRS_LRFD_CLOSED.DAT")
            return self.getFile("assests/templates/design/FORM_GRS_LRFD.DAT")
        elif design_code == Constants.ASCE_710_ASD:
            if enclosure == Constants.ENCLOSED_ROOF:
                return self.getFile("assests/templates/design/FORM_GRS_ASD_CLOSED.DAT")
            return self.getFile("assests/templates/design/FORM_GRS_ASD.DAT")

    def loadNonSeismicFiles(self, design_code, enclosure):
        print(design_code, enclosure)
        if design_code == Constants.ASCE_710_LRFD:
            if enclosure == Constants.ENCLOSED_ROOF:
                return Document(self.getFile("assests/templates/design/Gen-Desc_710_lrfd_closed.docx"))
            return Document(self.getFile("assests/templates/design/Gen-Desc_710_lrfd.docx"))
        elif design_code == Constants.ASCE_710_ASD:
            if enclosure == Constants.ENCLOSED_ROOF:
                return Document(self.getFile("assests/templates/design/Gen-Desc_710_asd_closed.docx"))
            return Document(self.getFile("assests/templates/design/Gen-Desc_710_asd.docx"))

    def loadSeismicFiles(self, design_code, enclosure):
        if design_code == Constants.ASCE_710_LRFD:
            if enclosure == Constants.ENCLOSED_ROOF:
                return Document(self.getFile("assests/templates/seismic/Seismic_Gen-Desc_710_lrfd_closed.docx"))
            return Document(self.getFile("assests/templates/seismic/Seismic_Gen-Desc_710_lrfd.docx"))
        elif design_code == Constants.ASCE_710_ASD:
            if enclosure == Constants.ENCLOSED_ROOF:
                return Document(self.getFile("assests/templates/seismic/Seismic_Gen-Desc_710_asd_closed.docx"))
            return Document(self.getFile("assests/templates/seismic/Seismic_Gen-Desc_710_asd.docx"))

    def getTable27_3_1(self):
        return json.load(open(self.getFile("assests/asce/table_27_3_1.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getWindDesignDefaults(self):
        return json.load(open(self.getFile("assests/templates/program/wind_design_defaults.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getWindCoeffiecients(self):
        return json.load(open(self.getFile("assests/templates/program/wind_coeffiecients.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getWindMapDefaults(self):
        return json.load(open(self.getFile("assests/templates/program/windmap_defaults.json"), encoding='utf8'), 
            object_pairs_hook=OrderedDict)

    def getInputFilePath(self):
        # if called with no arguments, call app data pick file from there
        path = None
        if (len(sys.argv) > 1):
            path = self.getFile(sys.argv[1])
            
        else:
            path = self.getFile(os.getenv('LOCALAPPDATA') + "\\Trevexs SSC\\data\\sample1.trsc")
        
        return path

    def getOutPutFilePath(self, name):
        try: 
            head = os.path.split(self.path)[0]
            file_name = "Gen-Desc-%s.DOC" %name
            return os.path.join(head, file_name)

        except AttributeError:
            # Messages.showError("There is no data to use to generate output file")
            print("An error occured in getOutPutFile method")

    def getRootOutPutPath(self, file_name):
        try: 
            head = os.path.split(self.path)[0]
            return os.path.join(head, file_name)

        except AttributeError:
            # Messages.showError("There is no data to use to generate output file")
            print("An error occured in getOutPutFile method")

    def getGeomFile(self):
        try: 
            head = os.path.split(self.path)[0]
            file_name = "GEOM1.DXF"
            return os.path.join(head, file_name)

        except AttributeError:
            # Messages.showError("There is no data to use to generate output file")
            print("Error: Cannot create geom1.dxf path")
            # print(error)


    def getLoadsFile(self):
        '''
        opens file, remmember to close file after file has
        opened
        '''
        try:
            head = os.path.split(self.path)[0]
            file_name = "loads.txt"
            path = os.path.join(head, file_name)
            return open(path, 'w+')
        except AttributeError:
            # Messages.showError("There is no data to use to generate output file")
            print("An error occured while creating loads file")

    def getFormGrmsFile(self):
        '''
        opens file, remmember to close file after file has
        opened
        '''
        try:
            head = os.path.split(self.path)[0]
            file_name = "FORM_GRS.DAT"
            path = os.path.join(head, file_name)
            return open(path, 'w+')
        except AttributeError:
            # Messages.showError("There is no data to use to generate output file")
            print("An error occured while creating loads file")
        
    def getWindMapImagePath(self):
        try: 
            head = os.path.split(self.path)[0]
            file_name = "WINDMAP.JPG"
            return os.path.join(head, file_name)

        except AttributeError:
            # Messages.showError("There is no data to use to generate output file")
            print("An error occured in getOutPutFile method")